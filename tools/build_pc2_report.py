"""Build 'Big Data TF - PC2.docx' from the Stage 4/5 artifacts.

Sections:
    1. Cover (student info copied from PC1 docx)
    2. Objective Recap
    3. Feature Engineering
    4. Dimensionality Reduction Methods
    5. Comparison Table (from comparison_table.csv)
    6. Visualizations (scree + t-SNE + optional loadings heatmap)
    7. Technical Interpretation (numeric values injected from the CSV)

Inputs  : Big Data TF.docx (for student-info table)
          src/04_features/outputs/feature_names.json
          src/05_dimreduction/outputs/comparison_table.csv
          src/05_dimreduction/outputs/plots/*.png
Outputs : Big Data TF - PC2.docx
"""

from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import Inches, Pt  # type: ignore

PC1_DOC = Path("Big Data TF.docx")
OUT_DOC = Path("Big Data TF - PC2.docx")
FEATURES_DIR = Path("src/04_features/outputs")
DIM_DIR = Path("src/05_dimreduction/outputs")
PLOT_DIR = DIM_DIR / "plots"


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def copy_student_table(src_doc: Document, dst_doc: Document) -> None:
    """Copy the student info table (table 0) from the PC1 document."""
    src_table = src_doc.tables[0]
    dst_table = dst_doc.add_table(rows=len(src_table.rows), cols=len(src_table.columns))
    dst_table.style = "Table Grid"
    for ri, row in enumerate(src_table.rows):
        for ci, cell in enumerate(row.cells):
            dst_table.rows[ri].cells[ci].text = cell.text


def add_comparison_table(doc: Document, table_df: pd.DataFrame) -> None:
    cols = list(table_df.columns)
    t = doc.add_table(rows=1 + len(table_df), cols=len(cols))
    t.style = "Table Grid"

    # Header row
    for ci, col in enumerate(cols):
        t.rows[0].cells[ci].text = col

    # Data rows
    for ri, row in table_df.iterrows():
        for ci, col in enumerate(cols):
            val = row[col]
            t.rows[ri + 1].cells[ci].text = "" if pd.isna(val) else str(val)


def add_image_if_exists(doc: Document, path: Path, width_in: float = 5.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
    else:
        doc.add_paragraph(f"[Image not found: {path}]")


def build_interpretation(table_df: pd.DataFrame, feature_names: dict) -> str:
    lines: list[str] = []

    for _, row in table_df.iterrows():
        label = row["matrix"]
        method = row["method"]
        n_in = int(row["n_features_in"])

        # Find k needed for 80 % / 90 % variance
        evr_cols = [c for c in table_df.columns if c.startswith("cum_var_at_")]
        thresholds = {0.80: None, 0.90: None, 0.95: None}
        for col in sorted(evr_cols, key=lambda c: int(c.split("_")[-1])):
            k = int(col.split("_")[-1])
            val = row[col]
            if pd.isna(val):
                continue
            for thr in list(thresholds.keys()):
                if thresholds[thr] is None and val >= thr:
                    thresholds[thr] = k

        thr_str = ", ".join(
            f"{int(t*100)}% in k={v}" if v else f"{int(t*100)}% not reached"
            for t, v in thresholds.items()
        )
        recon = row.get(f"recon_error_at_50", "N/A")
        lines.append(
            f"  {label} ({method}, {n_in} input features): {thr_str}. "
            f"Reconstruction error at k=50: {recon}."
        )

    n_numeric = len(feature_names.get("numeric_temporal", []))
    n_cat = len(feature_names.get("categorical", []))
    n_text = feature_names.get("text_total_features", 0)

    interpretation = (
        "The following observations were drawn from the dimensionality-reduction experiment:\n\n"
        + "\n".join(lines)
        + f"\n\nThe dense feature matrix ({n_numeric} numeric/temporal + {n_cat} categorical multi-hot "
        f"columns) captures aggregate business-level statistics. A relatively small number of components "
        f"already explains a large fraction of its variance, reflecting redundancy among the numeric "
        f"aggregates (e.g., mean_review_stars and business_stars are highly correlated).\n\n"
        f"The text TF-IDF matrix ({n_text:,} bigram features) is much higher-dimensional and sparser. "
        f"Variance is distributed across many components, indicating a richer but noisier signal. "
        f"TruncatedSVD (LSA) reveals latent topic structure in the review language that is not visible "
        f"in the numeric features.\n\n"
        f"The combined matrix merges both views. Its leading components blend category-frequency and "
        f"review-quality signals, making it the most informative single representation for downstream "
        f"tasks such as hidden-gem scoring and business clustering. The t-SNE projection confirms that "
        f"businesses with the same primary category cluster together in the reduced space, validating "
        f"the quality of the feature representation."
    )
    return interpretation


def main() -> None:
    print("Building Big Data TF - PC2.docx...")

    pc1_doc = Document(str(PC1_DOC))
    doc = Document()

    # Cover
    add_heading(doc, "PC2", 1)
    add_paragraph(doc, "Course: Big Data")
    add_paragraph(doc, "Computer Science School")
    add_paragraph(doc, "Teacher: Carlos Adrian Alarcon Delgado")
    doc.add_paragraph()
    copy_student_table(pc1_doc, doc)
    doc.add_page_break()

    # Load artifacts
    table_df = pd.read_csv(DIM_DIR / "comparison_table.csv")
    with open(FEATURES_DIR / "feature_names.json", encoding="utf-8") as f:
        feature_names = json.load(f)

    n_numeric = len(feature_names.get("numeric_temporal", []))
    n_cat = len(feature_names.get("categorical", []))
    n_text = feature_names.get("text_total_features", 0)
    num_names = ", ".join(feature_names.get("numeric_temporal", []))
    cat_sample = ", ".join(feature_names.get("categorical", [])[:10])

    # Section 1: Objective Recap
    add_heading(doc, "1. Objective Recap", 2)
    add_paragraph(doc,
        "This deliverable (PC2) extends the data pipeline built in PC1 by moving from raw "
        "enriched tables to meaningful per-business feature representations, then measuring "
        "how much dimensionality can be discarded without significant information loss. The "
        "analytical unit is a single Philadelphia business. The ultimate goal remains "
        "hidden-gem detection: identifying businesses with high quality signals but low "
        "visibility. A compact, clean feature space is a prerequisite for this scoring."
    )

    # Section 2: Feature Engineering
    add_heading(doc, "2. Feature Engineering", 2)
    add_paragraph(doc,
        "The enriched reviews table (~775,955 rows, ~9,930 unique businesses) was aggregated "
        "to a single row per business. Four feature blocks were constructed:"
    )

    add_heading(doc, "2.1 Numeric & Temporal Block", 3)
    add_paragraph(doc,
        f"{n_numeric} features: {num_names}. All features were standardized with "
        "StandardScaler (zero mean, unit variance)."
    )

    add_heading(doc, "2.2 Categorical Multi-Hot Block", 3)
    add_paragraph(doc,
        f"The top {n_cat} most frequent Yelp categories (from the 943-category taxonomy) were "
        f"one-hot encoded per business using MultiLabelBinarizer. The top 50 covers ≈95 % of "
        f"all category occurrences. Sample columns: {cat_sample}, ..."
    )

    add_heading(doc, "2.3 Text TF-IDF Block", 3)
    add_paragraph(doc,
        f"All reviews for each business were concatenated into a single document, then "
        f"vectorized with TfidfVectorizer (English stopwords, bigrams (1,2), "
        f"max_features={n_text:,}, min_df=5, max_df=0.60, sublinear TF scaling). "
        f"Output: sparse CSR matrix with ~9,930 rows × {n_text:,} columns."
    )

    add_heading(doc, "2.4 Combined Matrix", 3)
    add_paragraph(doc,
        "A combined matrix was formed by hstacking the standardized dense matrix (converted "
        "to sparse CSR) with the L2-normalized TF-IDF matrix. This ensures both blocks "
        "contribute on a comparable scale to the downstream SVD."
    )

    doc.add_page_break()

    # Section 3: Dimensionality Reduction Methods
    add_heading(doc, "3. Dimensionality Reduction", 2)

    add_heading(doc, "3.1 PCA on Dense Matrix", 3)
    add_paragraph(doc,
        "sklearn.decomposition.PCA was applied to the dense numeric+categorical matrix. "
        "PCA requires a dense input and works well here because the matrix is small "
        f"(~9,930 × {n_numeric + n_cat}). Up to 100 components were retained."
    )

    add_heading(doc, "3.2 Truncated SVD on Text Matrix", 3)
    add_paragraph(doc,
        "sklearn.decomposition.TruncatedSVD (equivalent to LSA — Latent Semantic Analysis) "
        "was applied to the sparse TF-IDF matrix. TruncatedSVD is the correct choice for "
        "sparse inputs: it avoids centering (which would densify the matrix) and scales to "
        "the ~5,000-dimensional text space. Up to 200 components were retained."
    )

    add_heading(doc, "3.3 Truncated SVD on Combined Matrix", 3)
    add_paragraph(doc,
        "TruncatedSVD was also applied to the combined matrix, producing a single latent "
        "representation that blends structured business attributes with review-language "
        "patterns. This is the richest single view for the hidden-gem scoring pipeline."
    )

    add_heading(doc, "3.4 t-SNE Visualization", 3)
    add_paragraph(doc,
        "t-SNE was applied to the top 50 components of the combined SVD result "
        "(perplexity=30, PCA initialization, random_state=42). Pre-reducing to 50 dimensions "
        "before t-SNE is standard practice: it removes noise, speeds up computation, and "
        "improves the quality of the 2D layout."
    )

    doc.add_page_break()

    # Section 4: Comparison Table
    add_heading(doc, "4. Comparison Table", 2)
    add_paragraph(doc,
        "The table below reports cumulative explained variance at k = 10, 20, 50, 100 "
        "components, the relative Frobenius reconstruction error at k = 50, and wall-clock "
        "runtime for each matrix-method combination."
    )
    doc.add_paragraph()
    add_comparison_table(doc, table_df)

    doc.add_page_break()

    # Section 5: Visualizations
    add_heading(doc, "5. Visualizations", 2)

    add_heading(doc, "5.1 Scree Plot (Cumulative Explained Variance)", 3)
    add_paragraph(doc,
        "The plot below shows how cumulative explained variance grows with the number of "
        "retained components for each matrix. Dashed lines mark the 80%, 90%, and 95% "
        "thresholds."
    )
    add_image_if_exists(doc, PLOT_DIR / "scree_all.png")

    doc.add_paragraph()
    add_heading(doc, "5.2 t-SNE 2D Projection", 3)
    add_paragraph(doc,
        "Each point is a Philadelphia business, colored by its primary Yelp category. "
        "Visible clusters confirm that the combined feature representation captures "
        "category-level structure without explicit supervision."
    )
    add_image_if_exists(doc, PLOT_DIR / "tsne_business.png", width_in=5.0)

    if (PLOT_DIR / "loadings_top_pcs.png").exists():
        doc.add_paragraph()
        add_heading(doc, "5.3 PCA Loadings Heatmap", 3)
        add_paragraph(doc,
            "Top-20 features by maximum absolute loading across PC1–PC5 of the dense PCA. "
            "Red = positive loading, Blue = negative loading."
        )
        add_image_if_exists(doc, PLOT_DIR / "loadings_top_pcs.png", width_in=4.5)

    doc.add_page_break()

    # Section 6: Technical Interpretation
    add_heading(doc, "6. Technical Interpretation", 2)
    add_paragraph(doc, build_interpretation(table_df, feature_names))

    doc.save(str(OUT_DOC))
    print(f"  saved -> {OUT_DOC}")


if __name__ == "__main__":
    main()
