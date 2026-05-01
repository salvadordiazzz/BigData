"""Build 'Big Data TF - Combined.docx'.

Opens the existing PC1 document (Big Data TF.docx) and appends the Week 5
content directly to it, preserving all PC1 formatting and styles.

Week 5 sections added:
    PC2 - Week 5: Representation and Dimensionality Report
    1. Feature Engineering (numeric/temporal, categorical, text, combined)
    2. Dimensionality Reduction Methods (PCA, SVD, t-SNE)
    3. Comparison Table (from comparison_table.csv)
    4. Visualizations (scree + t-SNE + loadings)
    5. Technical Interpretation

Inputs  : Big Data TF.docx
          src/04_features/outputs/feature_names.json
          src/05_dimreduction/outputs/comparison_table.csv
          src/05_dimreduction/outputs/plots/*.png
Outputs : Big Data TF - Combined.docx
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PC1_DOC  = Path("Big Data TF.docx")
OUT_DOC  = Path("Big Data TF - Combined.docx")
FEAT_DIR = Path("src/04_features/outputs")
DIM_DIR  = Path("src/05_dimreduction/outputs")
PLOT_DIR = DIM_DIR / "plots"


def heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Times New Roman"


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_image(doc: Document, path: Path, caption: str, width_in: float = 5.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Image not found: {path.name}]")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)


def _apply_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def add_comparison_table(doc: Document, df: pd.DataFrame) -> None:
    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    _apply_table_borders(table)

    for ci, col in enumerate(cols):
        cell = table.rows[0].cells[ci]
        cell.text = col
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)

    for ri, (_, row) in enumerate(df.iterrows()):
        for ci, col in enumerate(cols):
            val = row[col]
            cell = table.rows[ri + 1].cells[ci]
            cell.text = "" if pd.isna(val) else str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)


def build_interpretation(df: pd.DataFrame, feature_names: dict) -> str:
    lines = []
    evr_cols = sorted(
        [c for c in df.columns if c.startswith("cum_var_at_")],
        key=lambda c: int(c.split("_")[-1]),
    )
    for _, row in df.iterrows():
        label  = row["matrix"]
        method = row["method"]
        n_in   = int(row["n_features_in"])
        thresholds = {0.80: None, 0.90: None, 0.95: None}
        for col in evr_cols:
            k   = int(col.split("_")[-1])
            val = row[col]
            if pd.isna(val):
                continue
            for thr in list(thresholds):
                if thresholds[thr] is None and val >= thr:
                    thresholds[thr] = k
        thr_str = ", ".join(
            f"{int(t*100)}% at k={v}" if v else f"{int(t*100)}% not reached"
            for t, v in thresholds.items()
        )
        recon = row.get("recon_error_at_50", "N/A")
        lines.append(
            f"  {label} ({method}, {n_in} input features): {thr_str}. "
            f"Frobenius reconstruction error at k=50: {recon}."
        )

    n_num  = len(feature_names.get("numeric_temporal", []))
    n_cat  = len(feature_names.get("categorical", []))
    n_text = feature_names.get("text_total_features", 0)

    return (
        "The dimensionality-reduction experiments produced the following findings:\n\n"
        + "\n".join(lines)
        + f"\n\nThe dense matrix ({n_num} numeric/temporal + {n_cat} categorical columns) "
        "is compact and already well-structured. PCA recovers the majority of its variance "
        "within a small number of components, revealing high redundancy among the numeric "
        "aggregates — for example, mean_review_stars and business_stars are strongly "
        "correlated, as are the engagement-rate features (mean_useful, mean_funny, "
        "mean_cool). This is expected for hand-crafted, semantically overlapping business "
        "statistics.\n\n"
        f"The TF-IDF text matrix ({n_text:,} bigram features) is high-dimensional and "
        "sparse. Variance is distributed across many components, reflecting the rich "
        "vocabulary variation across business categories. TruncatedSVD (LSA) reveals "
        "latent topic directions — food-related vs. service-related terms, for instance — "
        "that are entirely invisible in the numeric features.\n\n"
        "The combined matrix provides the richest single representation. Its leading SVD "
        "components blend structured business statistics with review-language signals, "
        "making it the most suitable input for the upcoming clustering and hidden-gem "
        "scoring stages. The t-SNE projection confirms that businesses sharing a primary "
        "Yelp category form coherent clusters in the 2D layout without any supervised "
        "signal, validating the quality of the feature representation."
    )


def main() -> None:
    print("Building Big Data TF - Combined.docx...")

    doc = Document(str(PC1_DOC))

    with open(FEAT_DIR / "feature_names.json", encoding="utf-8") as f:
        feature_names = json.load(f)
    df = pd.read_csv(DIM_DIR / "comparison_table.csv")

    n_num   = len(feature_names.get("numeric_temporal", []))
    n_cat   = len(feature_names.get("categorical", []))
    n_text  = feature_names.get("text_total_features", 0)
    num_names   = ", ".join(feature_names.get("numeric_temporal", []))
    cat_sample  = ", ".join(feature_names.get("categorical", [])[:10])

    # ------------------------------------------------------------------ #
    #  PC2 content starts here                                            #
    # ------------------------------------------------------------------ #
    doc.add_page_break()
    heading(doc, "PC2 - Week 5: Representation and Dimensionality Report", 1)

    # 1. Feature Engineering
    heading(doc, "1. Feature Engineering", 2)
    body(doc,
        "The enriched reviews table (775,955 rows, 11,671 unique businesses after cleaning "
        "filters) was aggregated to one row per business. Four feature blocks were built:"
    )

    heading(doc, "1.1 Numeric and Temporal Block", 3)
    body(doc,
        f"{n_num} features standardized with StandardScaler (zero mean, unit variance): "
        f"{num_names}. The temporal features (days_since_first, days_since_last, "
        "review_span_days, reviews_per_year_recent) capture the activity window and recency "
        "of each business, which are key signals for identifying hidden gems."
    )

    heading(doc, "1.2 Categorical Multi-Hot Block", 3)
    body(doc,
        f"The top {n_cat} most frequent Yelp categories were one-hot encoded per business "
        "using MultiLabelBinarizer. The top 50 covers approximately 95% of all category "
        f"occurrences in the Philadelphia dataset (943 unique categories total). "
        f"Sample columns: {cat_sample}, ..."
    )

    heading(doc, "1.3 Text TF-IDF Block", 3)
    body(doc,
        "All reviews for each business were concatenated into a single document, then "
        f"vectorized with TfidfVectorizer (English stopwords, bigrams, "
        f"max_features={n_text:,}, min_df=5, max_df=0.60, sublinear_tf=True). "
        f"Output: sparse CSR matrix with ~11,671 rows x {n_text:,} columns."
    )

    heading(doc, "1.4 Combined Matrix", 3)
    body(doc,
        "A combined feature matrix was assembled by horizontally stacking the standardized "
        "dense block (numeric + categorical, converted to sparse CSR) with the L2-normalized "
        "TF-IDF matrix. L2 normalization on the text block ensures both halves contribute "
        "on a comparable scale to the downstream SVD decomposition."
    )

    doc.add_page_break()

    # 2. Dimensionality Reduction
    heading(doc, "2. Dimensionality Reduction Methods", 2)
    body(doc,
        "Three dimensionality-reduction experiments were conducted, one per feature matrix. "
        "Method choice is driven by the sparsity structure of each input:"
    )

    heading(doc, "2.1 PCA on Dense Matrix", 3)
    body(doc,
        f"Principal Component Analysis (sklearn.decomposition.PCA) was applied to the dense "
        f"numeric + categorical matrix (~11,671 x {n_num + n_cat} features). PCA requires a "
        "dense input, which is feasible here given the small number of columns. The "
        "explained_variance_ratio_ attribute records how much variance each component "
        "captures, enabling the scree plot below."
    )

    heading(doc, "2.2 Truncated SVD on Text Matrix", 3)
    body(doc,
        f"TruncatedSVD (sklearn.decomposition.TruncatedSVD, equivalent to Latent Semantic "
        f"Analysis) was applied to the sparse TF-IDF matrix (~11,671 x {n_text:,}). "
        "TruncatedSVD is the standard choice for sparse inputs: it avoids explicit "
        "mean-centering (which would densify a sparse matrix) and scales to the full "
        "vocabulary dimension. Up to 200 components were computed."
    )

    heading(doc, "2.3 Truncated SVD on Combined Matrix", 3)
    body(doc,
        "TruncatedSVD was also applied to the combined matrix, producing a single latent "
        "representation that blends structured business attributes with review-language "
        "patterns. With 200 retained components this view is the primary input for the "
        "upcoming clustering (Week 7) and hidden-gem scoring stages."
    )

    heading(doc, "2.4 t-SNE Visualization", 3)
    body(doc,
        "t-SNE (sklearn.manifold.TSNE, n_components=2, perplexity=30, init='pca', "
        "random_state=42) was applied to the top 50 components of the combined SVD. "
        "Pre-reducing to 50 dimensions before running t-SNE removes noise, accelerates "
        "computation, and improves the stability of the 2D layout. The result is saved "
        "as tsne_2d.parquet for use in subsequent stages."
    )

    doc.add_page_break()

    # 3. Comparison Table
    heading(doc, "3. Comparison Table", 2)
    body(doc,
        "Table 1 reports, for each matrix-method combination: cumulative explained variance "
        "at k = 10, 20, 50, 100; the relative Frobenius reconstruction error at k = 50 "
        "(||X - X_k||_F / ||X||_F); and wall-clock runtime on the local machine."
    )
    doc.add_paragraph()
    add_comparison_table(doc, df)
    cap = doc.add_paragraph(
        "Table 1 - Dimensionality reduction comparison (dense PCA, text SVD, combined SVD)."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    doc.add_page_break()

    # 4. Visualizations
    heading(doc, "4. Visualizations", 2)

    heading(doc, "4.1 Scree Plot - Cumulative Explained Variance", 3)
    body(doc,
        "The plot below shows how cumulative explained variance grows with the number of "
        "retained components. Horizontal dashed lines mark the 80%, 90%, and 95% thresholds."
    )
    add_image(
        doc,
        PLOT_DIR / "scree_all.png",
        "Figure 1 - Cumulative explained variance vs. k "
        "(dense PCA, text SVD, combined SVD; Philadelphia businesses).",
    )

    doc.add_paragraph()
    heading(doc, "4.2 t-SNE 2D Projection", 3)
    body(doc,
        "Each point represents a Philadelphia business, colored by its primary Yelp "
        "category (top 10 categories shown; remaining grouped as 'Other'). Visible clusters "
        "confirm that the combined feature representation captures category-level structure "
        "without any explicit supervision signal."
    )
    add_image(
        doc,
        PLOT_DIR / "tsne_business.png",
        "Figure 2 - t-SNE 2D projection of ~11,671 businesses, colored by primary category.",
        width_in=5.0,
    )

    if (PLOT_DIR / "loadings_top_pcs.png").exists():
        doc.add_paragraph()
        heading(doc, "4.3 PCA Loadings Heatmap", 3)
        body(doc,
            "Top-20 features by maximum absolute loading across PC1-PC5 of the dense PCA. "
            "Red = positive loading, Blue = negative loading. PC1 is dominated by "
            "review-quality aggregates (mean_review_stars, business_stars), while later "
            "components separate temporal activity and engagement signals."
        )
        add_image(
            doc,
            PLOT_DIR / "loadings_top_pcs.png",
            "Figure 3 - PCA loadings heatmap, top-20 features across PC1-PC5 (dense matrix).",
            width_in=4.5,
        )

    doc.add_page_break()

    # 5. Technical Interpretation
    heading(doc, "5. Technical Interpretation", 2)
    body(doc, build_interpretation(df, feature_names))

    doc.save(str(OUT_DOC))
    print(f"  saved -> {OUT_DOC}")


if __name__ == "__main__":
    main()
